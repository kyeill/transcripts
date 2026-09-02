"""Write a formatted .docx from aligned service blocks."""
from docx import Document
from docx.shared import Pt

from align import NEVER_SPOKEN_LABELS


def render(title, date_str, blocks, out_path):
    doc = Document()

    doc.add_heading(title, level=0)
    doc.add_paragraph(date_str).italic = True

    for block in blocks:
        if block.kind == "music":
            p = doc.add_paragraph()
            label = f"{block.label}"
            if block.title:
                label += f" — {block.title}"
            run = p.add_run(f"♪ {label}")
            run.italic = True
            run.font.size = Pt(11)
            continue

        heading = block.speaker or block.label
        p = doc.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        if block.speaker and block.speaker != block.label:
            sub = p.add_run(f"  ({block.label})")
            sub.italic = True
            sub.font.size = Pt(9)

        if block.text:
            doc.add_paragraph(block.text)
        elif block.label not in NEVER_SPOKEN_LABELS:
            # silence is the point of a silent meditation, so say nothing;
            # anywhere else an empty section is worth flagging
            doc.add_paragraph("[no speech captured]")

    doc.save(out_path)


if __name__ == "__main__":
    import sys

    from worship_guide import parse_worship_guide
    from align import align

    items = parse_worship_guide(sys.argv[1])
    blocks = align(items, [])
    render("Test", "smoke test", blocks, sys.argv[2])
    print("wrote", sys.argv[2])
